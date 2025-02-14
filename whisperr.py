import whisper
import streamlit as st
import os
from deep_translator import GoogleTranslator
from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
from transformers import pipeline
import pandas as pd
from keybert import KeyBERT
from gtts import gTTS
from docx import Document
import gtts.lang  # Import gtts.lang to access supported languages

# Page Configuration
st.set_page_config(page_title="EchoSense", page_icon="🎙️", layout="wide")
# Add a header for the app
st.title("🎙️ EchoSense")  # This will display the app name at the top
st.markdown("Welcome to EchoSense! This app transcribes audio, analyzes sentiment, and translates text.")

# Sidebar for Settings
st.sidebar.header("⚙️ Settings")
record_duration = st.sidebar.slider("Select recording duration (seconds)", min_value=1, max_value=300, value=5)
selected_language = st.sidebar.selectbox("Select translation language:", 
    {"english": "en", "persian": 'fa', "spanish": "es", "korean": "ko", "french": "fr", "german": "de", 
     "pashto": "ps", "chinese": "zh-CN", "urdu": "ur", "afrikaans": "af", "albanian": "sq", 
     "amharic": "am", "arabic": "ar", "armenian": "hy", "assamese": "as", "aymara": "ay", 
     "azerbaijani": "az", "bambara": "bm", "basque": "eu", "belarusian": "be", "bengali": "bn", 
     "bhojpuri": "bho", "bosnian": "bs", "bulgarian": "bg", "catalan": "ca", "cebuano": "ceb", 
     "chichewa": "ny", "chinese (simplified)": "zh-CN", "chinese (traditional)": "zh-TW", 
     "corsican": "co", "croatian": "hr", "czech": "cs", "danish": "da", "dhivehi": "dv", 
     "dogri": "doi", "dutch": "nl", "english": "en", "esperanto": "eo", "estonian": "et", 
     "ewe": "ee", "filipino": "tl", "finnish": "fi", "french": "fr", "frisian": "fy", 
     "galician": "gl", "georgian": "ka", "german": "de", "greek": "el", "guarani": "gn", 
     "gujarati": "gu", "haitian creole": "ht", "hausa": "ha", "hawaiian": "haw", "hebrew": "iw", 
     "hindi": "hi", "hmong": "hmn", "hungarian": "hu", "icelandic": "is", "igbo": "ig", 
     "ilocano": "ilo", "indonesian": "id", "irish": "ga", "italian": "it", "japanese": "ja", 
     "javanese": "jw", "kannada": "kn", "kazakh": "kk", "khmer": "km", "kinyarwanda": "rw", 
     "konkani": "gom", "korean": "ko", "krio": "kri", "kurdish (kurmanji)": "ku", 
     "kurdish (sorani)": "ckb", "kyrgyz": "ky", "lao": "lo", "latin": "la", "latvian": "lv", 
     "lingala": "ln", "lithuanian": "lt", "luganda": "lg", "luxembourgish": "lb", 
     "macedonian": "mk", "maithili": "mai", "malagasy": "mg", "malay": "ms", "malayalam": "ml", 
     " maltese": "mt", "maori": "mi", "marathi": "mr", "meiteilon (manipuri)": "mni-Mtei", 
     "mizo": "lus", "mongolian": "mn", "myanmar": "my", "nepali": "ne", "norwegian": "no", 
     "odia (oriya)": "or", "oromo": "om", "pashto": "ps", "persian": "fa", "polish": "pl", 
     "portuguese": "pt", "punjabi": "pa", "quechua": "qu", "romanian": "ro", "russian": "ru", 
     "samoan": "sm", "sanskrit": "sa", "scots gaelic": "gd", "sepedi": "nso", "serbian": "sr", 
     "sesotho": "st", "shona": "sn", "sindhi": "sd", "sinhala": "si", "slovak": "sk", 
     "slovenian": "sl", "somali": "so", "spanish": "es", "sundanese": "su", "swahili": "sw", 
     "swedish": "sv", "tajik": "tg", "tamil": "ta", "tatar": "tt", "telugu": "te", "thai": "th", 
     "tigrinya": "ti", "tsonga": "ts", "turkish": "tr", "turkmen": "tk", "twi": "ak", 
     "ukrainian": "uk", "urdu": "ur", "uyghur": "ug", "uzbek": "uz", "vietnamese": "vi", 
     "welsh": "cy", "xhosa": "xh", "yiddish": "yi", "yoruba": "yo", "zulu": "zu"})

# Initialize Session State
for key in ["transcribed_text", "detected_lang", "translated_text", "sentiment", "emotion", "keywords"]:
    if key not in st.session_state:
        st.session_state[key] = None

@st.cache_resource
def load_whisper_model():
    return whisper.load_model("large")

# Load Whisper Model
model = load_whisper_model()
st.sidebar.success("✅ Whisper model loaded!")

# Upload or Record Audio
col1, col2 = st.columns(2)

# Upload Audio File
with col1:
    audio_file = st.file_uploader("📂 Upload an audio file", type=["mp3", "wav", "m4a"])

# Record Audio
with col2:
    if st.button("🎙️ Start Recording"):
        with st.spinner("Recording..."):
            sample_rate = 44100
            audio_data = sd.rec(int(record_duration * sample_rate), samplerate=sample_rate, channels=2, dtype=np.int16)
            sd.wait()
            temp_audio_path = "recorded_audio.wav"
            wav.write(temp_audio_path, sample_rate, audio_data)
            st.success("✅ Recording complete!")
            st.audio(temp_audio_path, format="audio/wav")

# Determine Audio Path
audio_path = None
if audio_file:
    temp_audio_path = "uploaded_audio.wav"
    with open(temp_audio_path, "wb") as f:
        f.write(audio_file.read())  # Save uploaded file locally
    audio_path = temp_audio_path
    st.audio(temp_audio_path, format="audio/wav")
elif os.path.exists("recorded_audio.wav"):
    audio_path = "recorded_audio.wav"
if audio_path:
    st.download_button("📥 Download Audio", data=open(audio_path, "rb").read(), file_name=os.path.basename(audio_path), mime="audio/wav")

# Transcribe Audio
if audio_path and st.button("📝 Transcribe Audio"):
    with st.spinner("Transcribing..."):
        transcription = model.transcribe(audio_path, language=None)
        text = transcription["text"].strip()
        lang = transcription.get("language", "Unknown")

        # Sentiment Analysis
        vader_analyzer = SentimentIntensityAnalyzer()
        compound = vader_analyzer.polarity_scores(text)["compound"]
        sentiment = "😊 Positive" if compound >= 0.3 else "😠 Negative" if compound <= -0.3 else "😐 Neutral"

        # Emotion Analysis
        emotion_classifier = pipeline("text-classification", model="j-hartmann/emotion-english-distilroberta-base", return_all_scores=True)
        emotions = emotion_classifier(text)[0]
        top_emotion = max(emotions, key=lambda x: x['score'])

        # Keyword Extraction
        kw_model = KeyBERT()
        keywords = kw_model.extract_keywords (text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=5)

        # Store Results in Session State
        st.session_state.update({
            "transcribed_text": text,
            "detected_lang": lang,
            "sentiment": sentiment,
            "emotion": f"{top_emotion['label']} ({top_emotion['score']:.2f})",
            "keywords": ", ".join([kw[0] for kw in keywords])
        })

        st.success(f"✅ Transcription complete! (Language: {lang.upper()})")

# Display Results
if st.session_state["transcribed_text"]:
    st.markdown(f"### **🗣️ Language:** `{st.session_state['detected_lang'].upper()}`")
    st.markdown(f"### **📜 Transcription:**\n\n*{st.session_state['transcribed_text']}*")
    st.markdown(f"### **📊 Sentiment:** {st.session_state['sentiment']}")
    st.markdown(f"### **😃 Emotion:** {st.session_state['emotion']}")
    st.markdown(f"### **🔑 Keywords:** {st.session_state['keywords']}")

# Translate Text
if st.session_state["transcribed_text"] and st.button("🌐 Translate Text"):
    translated_text = GoogleTranslator(source="auto", target=selected_language).translate(st.session_state["transcribed_text"])
    st.session_state["translated_text"] = translated_text
    st.success(f"✅ Translated to {selected_language}!")

if st.session_state["translated_text"]:
    st.markdown(f"### **📝 Translated Text ({selected_language}):**\n\n*{st.session_state['translated_text']}*")

# Language mapping for gTTS
language_mapping = {
    "english": "en",
    "persian": "fa",
    "spanish": "es",
    "korean": "ko",
    "french": "fr",
    "german": "de",
    "pashto": "ps",
    "chinese": "zh-CN",
    "urdu": "ur",
    "afrikaans": "af",
    "albanian": "sq",
    "amharic": "am",
    "arabic": "ar",
    "armenian": "hy",
    "assamese": "as",
    "aymara": "ay",
    "azerbaijani": "az",
    "bambara": "bm",
    "basque": "eu",
    "belarusian": "be",
    "bengali": "bn",
    "bhojpuri": "bho",
    "bosnian": "bs",
    "bulgarian": "bg",
    "catalan": "ca",
    "cebuano": "ceb",
    "chichewa": "ny",
    "corsican": "co",
    "croatian": "hr",
    "czech": "cs",
    "danish": "da",
    "dhivehi": "dv",
    "dogri": "doi",
    "dutch": "nl",
    "esperanto": "eo",
    "estonian": "et",
    "ewe": "ee",
    "filipino": "tl",
    "finnish": "fi",
    "frisian": "fy",
    "galician": "gl",
    "georgian": "ka",
    "greek": "el",
    "guarani": "gn",
    "gujarati": "gu",
    "haitian creole": "ht",
    "hausa": "ha",
    "hawaiian": "haw",
    "hebrew": "iw",
    "hindi": "hi",
    "hmong": "hmn",
    "hungarian": "hu",
    "icelandic": "is",
    "igbo": "ig",
    "ilocano": "ilo",
    "indonesian": "id",
    "irish": "ga",
    "italian": "it",
    "japanese": "ja",
    "javanese": "jw",
    "kannada": "kn",
    "kazakh": "kk",
    "khmer": "km",
    "kinyarwanda": "rw",
    "konkani": "gom",
    "kurdish (kurmanji)": "ku",
    "kurdish (sorani)": "ckb",
    "kyrgyz": "ky",
    "lao": "lo",
    "latin": "la",
    "latvian": "lv",
    "lingala": "ln",
    "lithuanian": "lt",
    "luganda": "lg",
    "luxembourgish": "lb",
    "macedonian": "mk",
    "maithili": "mai",
    "malagasy": "mg",
    "malay": "ms",
    "malayalam": "ml",
    "maltese": "mt",
    "maori": "mi",
    "marathi": "mr",
    "meiteilon (manipuri)": "mni-Mtei",
    "mizo": "lus",
    "mongolian": "mn",
    " myanmar": "my",
    "nepali": "ne",
    "norwegian": "no",
    "odia (oriya)": "or",
    "oromo": "om",
    "polish": "pl",
    "portuguese": "pt",
    "punjabi": "pa",
    "quechua": "qu",
    "romanian": "ro",
    "russian": "ru",
    "samoan": "sm",
    "sanskrit": "sa",
    "scots gaelic": "gd",
    "sepedi": "nso",
    "serbian": "sr",
    "sesotho": "st",
    "shona": "sn",
    "sindhi": "sd",
    "sinhala": "si",
    "slovak": "sk",
    "slovenian": "sl",
    "somali": "so",
    "sundanese": "su",
    "swahili": "sw",
    "swedish": "sv",
    "tajik": "tg",
    "tamil": "ta",
    "tatar": "tt",
    "telugu": "te",
    "thai": "th",
    "tigrinya": "ti",
    "tsonga": "ts",
    "turkish": "tr",
    "turkmen": "tk",
    "twi": "ak",
    "ukrainian": "uk",
    "uyghur": "ug",
    "uzbek": "uz",
    "vietnamese": "vi",
    "welsh": "cy",
    "xhosa": "xh",
    "yiddish": "yi",
    "yoruba": "yo",
    "zulu": "zu"
}

# Play Translated Audio
if st.session_state["translated_text"] and st.button("🔊 Play Translated Audio", key="play_translated_audio"):
    gtts_lang_code = language_mapping.get(selected_language)
    if gtts_lang_code:
        tts = gTTS(text=st.session_state["translated_text"], lang=gtts_lang_code, slow=False)
        audio_file_path = "translated_audio.mp3"
        tts.save(audio_file_path)
        
        # Play the audio file
        st.audio(audio_file_path, format="audio/mp3")
        st.success("✅ Playing the translated audio!")
    else:
        st.warning(f"⚠️ The selected language '{selected_language}' is not supported for speech synthesis.")

# Download Transcription & Translation
if st.session_state["transcribed_text"]:
    export_format = st.selectbox("Select export format:", ["Text File (.txt)", "CSV File (.csv)", "DOCX File (.docx)"])
    
    if st.button("💾 Download Transcription"):
        if export_format == "Text File (.txt)":
            file_content = f"Transcription:\n{st.session_state['transcribed_text']}\nLanguage: {st.session_state['detected_lang']}\nSentiment: {st.session_state['sentiment']}\nEmotion: {st.session_state['emotion']}\nKeywords: {st.session_state['keywords']}\nTranslated Text ({selected_language}): {st.session_state['translated_text']}"
            st.download_button("📥 Download TXT", file_content, "transcription.txt", "text/plain")
        elif export_format == "CSV File (.csv)":
            df = pd.DataFrame([{ 
                "Transcription": st.session_state['transcribed_text'], 
                "Language": st.session_state['detected_lang'], 
                "Sentiment": st.session_state['sentiment'], 
                "Emotion": st.session_state['emotion'], 
                "Keywords": st.session_state['keywords'], 
                "Translated Text": st.session_state['translated_text'] 
            }])
            st.download_button("📥 Download CSV", df.to_csv(index=False).encode(), "transcription.csv", "text/csv")
        elif export_format == "DOCX File (.docx)":
            doc = Document()
            doc.add_heading('Transcription', level=1)
            doc.add_paragraph(f"Transcription: {st.session_state['transcribed_text']}")
            doc.add_paragraph(f"Language: {st.session_state['detected_lang']}")
            doc.add_paragraph(f"Sentiment: {st.session_state['sentiment']}")
            doc.add_paragraph(f"Emotion: {st.session_state['emotion']}")
            doc.add_paragraph(f"Keywords: {st.session_state['keywords']}")
            doc.add_paragraph(f"Translated Text ({selected_language}): {st.session_state['translated_text']}")
            doc_path = "transcription.docx"
            doc.save(doc_path)
            with open(doc_path, "rb") as f:
                st.download_button("📥 Download DOCX", f, "transcription.docx", "application/vnd.openxmlformats-officedocument.word processingml.document")

# Clear Session State
if st.button("🗑️ Clear Session"):
    for key in st.session_state.keys():
        del st.session_state[key]
    st.success("✅ Session cleared!")

# Footer
st.markdown("---")